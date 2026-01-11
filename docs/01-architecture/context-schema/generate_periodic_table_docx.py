#!/usr/bin/env python3
"""Generate a DOCX document with chapters for each AI Periodic Table element."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import json
from pathlib import Path

# Load the JSON data
json_path = Path(__file__).parent / "ai-periodic-table.json"
with open(json_path) as f:
    data = json.load(f)

# Create document
doc = Document()

# Set up styles
title_style = doc.styles['Title']
title_style.font.size = Pt(28)
title_style.font.bold = True

heading1_style = doc.styles['Heading 1']
heading1_style.font.size = Pt(18)
heading1_style.font.bold = True

heading2_style = doc.styles['Heading 2']
heading2_style.font.size = Pt(14)
heading2_style.font.bold = True

# Title page
title = doc.add_heading(data['title'], 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph(data['subtitle'])
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph(f"Updated: {data['updated']}")
doc.add_paragraph(f"Framework Attribution: {data['framework_attribution']}")

# Add page break
doc.add_page_break()

# Table of Contents placeholder
doc.add_heading("Table of Contents", level=1)
doc.add_paragraph("See document navigation for complete chapter listing.")
doc.add_page_break()

# Executive Summary
doc.add_heading("Executive Summary", level=1)
stats = data['statistics']
summary_para = doc.add_paragraph()
summary_para.add_run("Competitive Position:\n\n").bold = True
summary_para.add_run(f"• Strong: {stats['strong']} elements\n")
summary_para.add_run(f"• Unique Differentiators: {stats['unique']} elements\n")
summary_para.add_run(f"• Emerging: {stats['emerging']} elements\n")
summary_para.add_run(f"• Gaps: {stats['gaps']} elements\n\n")
summary_para.add_run(
    "openContextSchema demonstrates Strong implementation in 80% of AI Periodic Table elements. "
    "The Context Series (Ea, Kv, Gk) represents the core industrial differentiator for enterprise AI."
)
doc.add_page_break()

# Legend
doc.add_heading("Legend", level=1)
for status, desc in data['legend'].items():
    p = doc.add_paragraph()
    p.add_run(f"{status.upper()}: ").bold = True
    p.add_run(desc)
doc.add_page_break()

# Status colors/labels for reference
STATUS_LABELS = {
    'strong': 'STRONG',
    'emerging': 'EMERGING',
    'gap': 'GAP',
    'unique': 'UNIQUE DIFFERENTIATOR'
}

def add_element_chapter(element_data, chapter_num):
    """Add a chapter for a single element."""
    symbol = element_data['symbol']
    name = element_data['name']
    
    # Chapter heading
    doc.add_heading(f"Chapter {chapter_num}: [{symbol}] {name}", level=1)
    
    # Quick reference box
    p = doc.add_paragraph()
    p.add_run("Quick Reference\n").bold = True
    p.add_run(f"Symbol: ").bold = True
    p.add_run(f"{symbol}\n")
    p.add_run(f"Classification: ").bold = True
    p.add_run(f"{element_data['classification']}\n")
    p.add_run(f"Status: ").bold = True
    p.add_run(f"{STATUS_LABELS.get(element_data['status'], element_data['status'].upper())}\n")
    p.add_run(f"Component: ").bold = True
    p.add_run(f"{element_data['component']}\n")
    
    # Grid position for main elements
    if 'row' in element_data:
        p.add_run(f"Grid Position: ").bold = True
        p.add_run(f"{element_data['row']}/{element_data['column']}\n")
    
    doc.add_paragraph()
    
    # Definition
    doc.add_heading("Definition", level=2)
    doc.add_paragraph(element_data['definition'])
    
    # Industrial Value
    doc.add_heading("Industrial Value", level=2)
    doc.add_paragraph(element_data['industrialValue'])
    
    # Reactivity & Bonds
    doc.add_heading("Reactivity & Bonds", level=2)
    if 'reactivity' in element_data:
        doc.add_paragraph(element_data['reactivity'])
    
    if 'bondsWith' in element_data:
        p = doc.add_paragraph()
        p.add_run("Bonds with: ").bold = True
        p.add_run(", ".join(element_data['bondsWith']))
    
    # Additional relationships for Context Series
    if 'catalyzedBy' in element_data:
        p = doc.add_paragraph()
        p.add_run("Catalyzed by: ").bold = True
        p.add_run(", ".join(element_data['catalyzedBy']))
    
    if 'stabilizedBy' in element_data:
        p = doc.add_paragraph()
        p.add_run("Stabilized by: ").bold = True
        p.add_run(", ".join(element_data['stabilizedBy']))
    
    if 'enhances' in element_data:
        p = doc.add_paragraph()
        p.add_run("Enhances: ").bold = True
        p.add_run(", ".join(element_data['enhances']))
    
    if 'prerequisiteFor' in element_data:
        p = doc.add_paragraph()
        p.add_run("Prerequisite for: ").bold = True
        p.add_run(", ".join(element_data['prerequisiteFor']))
    
    if 'outputOf' in element_data:
        p = doc.add_paragraph()
        p.add_run("Output of: ").bold = True
        p.add_run(", ".join(element_data['outputOf']))
    
    if 'triggers' in element_data:
        p = doc.add_paragraph()
        p.add_run("Triggers: ").bold = True
        p.add_run(", ".join(element_data['triggers']))
    
    if 'opposes' in element_data:
        p = doc.add_paragraph()
        p.add_run("Opposes: ").bold = True
        p.add_run(", ".join(element_data['opposes']))
    
    if 'overcomeBy' in element_data:
        p = doc.add_paragraph()
        p.add_run("Overcome by: ").bold = True
        p.add_run(", ".join(element_data['overcomeBy']))
    
    doc.add_page_break()

# Part 1: Main Grid Elements
doc.add_heading("Part I: Main Grid Elements", level=1)
doc.add_paragraph(
    "The main grid consists of elements organized by row (capability tier) and column (functional category). "
    "This mirrors the structure of the periodic table, where position indicates properties and relationships."
)
doc.add_page_break()

chapter_num = 1

# Order elements by row then column for logical progression
row_order = ['R1', 'R2', 'R3', 'R4']
col_order = ['C1', 'C2', 'C3', 'C4', 'C5']

# Create ordered list of elements
ordered_elements = []
for row in row_order:
    for col in col_order:
        for sym, elem in data['elements'].items():
            if elem.get('row') == row and elem.get('column') == col:
                ordered_elements.append(elem)
                break

for elem in ordered_elements:
    add_element_chapter(elem, chapter_num)
    chapter_num += 1

# Part 2: Context Series
doc.add_heading("Part II: The Context Series — Industrial AI Elements", level=1)
doc.add_paragraph(
    "The Context Series represents openContextSchema's core industrial differentiators. "
    "Like the lanthanides in the periodic table, these elements form a special series that "
    "underlies and enables the main grid functionality. These elements are specifically designed "
    "to overcome the 'Memory Wall' limitation of traditional LLM approaches."
)
doc.add_page_break()

# Context Series elements in logical order
context_order = ['Ec', 'Ea', 'Kv', 'Mp', 'Vr', 'Mw']
for sym in context_order:
    if sym in data['contextSeries']['elements']:
        add_element_chapter(data['contextSeries']['elements'][sym], chapter_num)
        chapter_num += 1

# Appendix: Element Relationships Diagram
doc.add_heading("Appendix A: Element Relationships", level=1)
doc.add_paragraph(
    "For an interactive visualization of element relationships and bonds, "
    "please refer to the interactive HTML version: ai-periodic-table-interactive.html"
)

# Save the document
output_path = Path(__file__).parent / "ai-periodic-table-guide.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
